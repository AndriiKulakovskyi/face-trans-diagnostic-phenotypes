"""Render ``docs/face_stratification/article.md`` as a detailed
scientific technical report in Word (``.docx``) format.

This script is a deliberately thin wrapper around *pandoc* (3.x).
It does **not** paraphrase, reshape or shorten the canonical
article markdown — the markdown file under
``docs/face_stratification/article.md`` is the single source of
truth for the technical report, and the Word version is
reproduced from it byte-for-byte (modulo pandoc formatting).

Design notes
------------

1. **No journal-style condensation.** An earlier version of this
   script hand-wrote a ~5-section journal abstract in Python
   strings. That approach drifted out of sync with the canonical
   markdown and was abandoned. The new script renders *all* 11
   sections, every table, every figure reference, the full
   appendices, and the clinical cluster-naming tables.

2. **Pandoc as the renderer.** Pandoc 3.8 handles GitHub-flavoured
   markdown tables, ATX headings, inline math and image embedding
   out-of-the-box and produces a well-formed .docx with a proper
   document outline, so downstream tooling (Word, LibreOffice,
   python-docx) can navigate it. We invoke it via ``subprocess``
   rather than ``pypandoc`` to avoid an extra Python dependency.

3. **Relative image paths.** The markdown references figures with
   paths like
   ``../../output/stratification/publication_figures/fig01_graph_structure.png``
   that are relative to ``docs/face_stratification/``. We run
   pandoc with ``cwd`` set to that directory so those paths
   resolve correctly without rewriting the markdown.

4. **Post-processing via python-docx.** After pandoc writes the
   .docx, we open it with ``python-docx`` to:

   * set document properties (title, author, subject, category);
   * add a running-head footer with a page-number field.

   The post-processing is best-effort: if it fails, the raw
   pandoc .docx is still a valid, complete technical report.

Usage
-----

::

    python scripts/generate_article_docx.py
    python scripts/generate_article_docx.py --output path/to/out.docx
    python scripts/generate_article_docx.py --no-toc --no-number-sections

"""

from __future__ import annotations

import argparse
import logging
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

logger = logging.getLogger("generate_article_docx")


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------


REPO_ROOT = Path(__file__).resolve().parent.parent
ARTICLE_MD = REPO_ROOT / "docs" / "face_stratification" / "article.md"
DEFAULT_OUTPUT = (
    REPO_ROOT
    / "output"
    / "stratification"
    / "transdiagnostic_stratification_article.docx"
)


# ---------------------------------------------------------------------------
# Configuration dataclass
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class RenderConfig:
    """Options controlling how the markdown is rendered to .docx."""

    article_md: Path
    output: Path
    include_toc: bool = True
    toc_depth: int = 3
    # The source markdown already carries manual section numbers
    # ("1. Introduction", "2. Related work", ...). Turning on pandoc's
    # --number-sections produces ugly double numbering like
    # "1.2   1. Introduction", so we leave it off by default and expose
    # an opt-in CLI flag for callers who explicitly want it.
    number_sections: bool = False
    standalone: bool = True
    highlight_style: str = "tango"
    reference_doc: Path | None = None

    def pandoc_args(self) -> list[str]:
        """Build the pandoc command-line arguments for this config."""

        # Pandoc is invoked with cwd = article_md.parent, so we pass the
        # basename here. That makes image paths like
        # ``../../output/stratification/publication_figures/fig*.png``
        # resolve correctly without rewriting the markdown file.
        args: list[str] = [
            self.article_md.name,
            "-o",
            str(self.output),
            "--from=gfm+tex_math_dollars+yaml_metadata_block",
            "--to=docx",
        ]
        if self.standalone:
            args.append("--standalone")
        if self.include_toc:
            args.append("--toc")
            args.append(f"--toc-depth={self.toc_depth}")
        if self.number_sections:
            args.append("--number-sections")
        if self.highlight_style:
            args.append(f"--highlight-style={self.highlight_style}")
        if self.reference_doc is not None and self.reference_doc.exists():
            args.append(f"--reference-doc={self.reference_doc}")
        return args


# ---------------------------------------------------------------------------
# Pandoc invocation
# ---------------------------------------------------------------------------


def _find_pandoc() -> str:
    """Locate the pandoc binary, preferring the conda-installed copy."""

    for candidate in ("/opt/anaconda3/bin/pandoc", "pandoc"):
        resolved = shutil.which(candidate)
        if resolved:
            return resolved
    raise RuntimeError(
        "pandoc is not installed or not on PATH. Install pandoc 3.x "
        "(https://pandoc.org/installing.html) before running this script."
    )


def _pandoc_version(pandoc: str) -> str:
    out = subprocess.run(
        [pandoc, "--version"], check=True, capture_output=True, text=True
    ).stdout.splitlines()[0]
    return out.strip()


def render_with_pandoc(config: RenderConfig) -> None:
    """Render ``config.article_md`` to ``config.output`` via pandoc."""

    if not config.article_md.exists():
        raise FileNotFoundError(f"article markdown not found: {config.article_md}")

    config.output.parent.mkdir(parents=True, exist_ok=True)

    pandoc = _find_pandoc()
    logger.info("Using %s (%s)", pandoc, _pandoc_version(pandoc))
    logger.info("Rendering %s", config.article_md)
    logger.info("Output    %s", config.output)

    cmd = [pandoc, *config.pandoc_args()]
    logger.debug("pandoc command: %s", " ".join(cmd))

    try:
        subprocess.run(
            cmd,
            check=True,
            cwd=config.article_md.parent,
            capture_output=True,
            text=True,
        )
    except subprocess.CalledProcessError as err:  # pragma: no cover - diagnostic
        logger.error("pandoc failed with exit code %s", err.returncode)
        if err.stdout:
            logger.error("stdout:\n%s", err.stdout)
        if err.stderr:
            logger.error("stderr:\n%s", err.stderr)
        raise


# ---------------------------------------------------------------------------
# Post-processing (optional, best-effort)
# ---------------------------------------------------------------------------


_DOC_METADATA = {
    "title": (
        "Data-driven transdiagnostic stratification of 11,014 psychiatric "
        "patients from the FACE cohort"
    ),
    "subject": "Precision psychiatry — graph representation learning — FACE cohort",
    "author": "Andrii Kulakovskyi, and the FACE Consortium",
    "category": "Technical report",
    "comments": (
        "Rendered from docs/face_stratification/article.md by "
        "scripts/generate_article_docx.py. Every quantitative claim is "
        "traceable to the stage documents under docs/face_stratification/ "
        "and the JSON summaries under output/stratification/."
    ),
}


def _add_page_number_field(paragraph) -> None:  # pragma: no cover - best effort
    """Append a Word PAGE field to the given paragraph."""

    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception as exc:
        logger.debug("python-docx internals unavailable: %s", exc)
        return

    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    for node in (fld_char_begin, instr_text, fld_char_separate, fld_char_end):
        run._r.append(node)


def postprocess_docx(path: Path) -> None:
    """Apply best-effort cosmetic post-processing to the pandoc output."""

    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "python-docx not available (%s); skipping post-processing", exc
        )
        return

    try:
        doc = Document(str(path))
    except Exception as exc:  # pragma: no cover
        logger.warning("could not open %s with python-docx: %s", path, exc)
        return

    # Document properties ---------------------------------------------------
    try:
        props = doc.core_properties
        props.title = _DOC_METADATA["title"]
        props.author = _DOC_METADATA["author"]
        props.subject = _DOC_METADATA["subject"]
        props.category = _DOC_METADATA["category"]
        props.comments = _DOC_METADATA["comments"]
    except Exception as exc:  # pragma: no cover
        logger.debug("failed to set core properties: %s", exc)

    # Running-head footer with page number ---------------------------------
    try:
        section = doc.sections[0]
        footer = section.footer
        for para in list(footer.paragraphs):
            para.clear()
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(
            "FACE transdiagnostic stratification — technical report · page "
        )
        _add_page_number_field(para)
    except Exception as exc:  # pragma: no cover
        logger.debug("failed to set footer: %s", exc)

    try:
        doc.save(str(path))
    except Exception as exc:  # pragma: no cover
        logger.warning("failed to save post-processed docx: %s", exc)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _parse_args(argv: Iterable[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Render the FACE transdiagnostic stratification article "
            "(docs/face_stratification/article.md) as a detailed .docx "
            "technical report."
        )
    )
    parser.add_argument(
        "--article",
        type=Path,
        default=ARTICLE_MD,
        help=f"Path to the source article markdown (default: {ARTICLE_MD}).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Path to the output .docx file (default: {DEFAULT_OUTPUT}).",
    )
    parser.add_argument(
        "--no-toc",
        dest="include_toc",
        action="store_false",
        help="Do not emit a table of contents.",
    )
    parser.add_argument(
        "--toc-depth",
        type=int,
        default=3,
        help="Table-of-contents depth (default: 3).",
    )
    parser.add_argument(
        "--number-sections",
        dest="number_sections",
        action="store_true",
        default=False,
        help=(
            "Prepend pandoc auto section numbers (1, 1.1, ...). Off by "
            "default because the source markdown already carries manual "
            "section numbers."
        ),
    )
    parser.add_argument(
        "--reference-doc",
        type=Path,
        default=None,
        help=(
            "Optional pandoc reference doc (.docx) whose styles will be "
            "used as the base for the output document."
        ),
    )
    parser.add_argument(
        "--no-postprocess",
        dest="postprocess",
        action="store_false",
        help="Skip python-docx cosmetic post-processing.",
    )
    parser.add_argument(
        "--verbose",
        "-v",
        action="store_true",
        help="Enable verbose logging.",
    )
    return parser.parse_args(list(argv) if argv is not None else None)


def main(argv: Iterable[str] | None = None) -> int:
    args = _parse_args(argv)
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )

    config = RenderConfig(
        article_md=args.article.resolve(),
        output=args.output.resolve(),
        include_toc=args.include_toc,
        toc_depth=args.toc_depth,
        number_sections=args.number_sections,
        reference_doc=args.reference_doc.resolve() if args.reference_doc else None,
    )

    try:
        render_with_pandoc(config)
    except Exception as exc:
        logger.error("pandoc render failed: %s", exc)
        return 1

    if args.postprocess:
        postprocess_docx(config.output)

    size_kb = config.output.stat().st_size / 1024.0
    logger.info("Wrote %s (%.1f KB)", config.output, size_kb)
    logger.info(
        "Done. Open the file in Word and verify: title, ToC, 11 sections, "
        "§7.0 cluster naming table, §7.3 six cluster subsections, §10.4, "
        "§10.5, Appendix A, Appendix B, and all figures."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
